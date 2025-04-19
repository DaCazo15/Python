import os, locale
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, timedelta
from pdf import word_to_pdf

def fecha(dia):
    locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
    fecha_actual = datetime.now()
    dias_semana = ["domingo", "lunes", "martes", "miércoles", "jueves", "viernes", "sábado"]
    dia_sistema = fecha_actual.strftime("%A").lower()
    dia_requerido = (dias_semana.index(dia) - dias_semana.index(dia_sistema)) * -1
    fecha_deseada = fecha_actual - timedelta(days=dia_requerido)
    return dias_semana, fecha_deseada, fecha_actual

def mp(nombre, escuela, dia):

    documento = Document()
    if dia == "sabado": 
        dia = "sábado"
    elif dia == "miercoles":
        dia = "miércoles"
    dias, dia_numero, fecha_deseada = fecha(dia)
    fecha_formateada = fecha_deseada.strftime(f"{dia.capitalize()} {dia_numero.strftime("%d")} de {"%B".capitalize()} de %Y")
    print(fecha_formateada)
    if os.path.exists("img/logo.png"):
        documento.add_picture("img/logo.png", width=Inches(2))
        linea = documento.paragraphs[-1]
        linea.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    escuelas_formateadas = ', '.join([esc.capitalize() for esc in escuela])
    documento.add_heading(f"Multa por Planificacion", 0) 
    documento.add_paragraph(fecha_deseada.strftime("%A %d de %Y"))
    documento.add_paragraph("DEPARTAMENTO DE ESCUELAS")
    documento.add_paragraph("Encargado: Ing. Daniel Cazorla")
    documento.add_paragraph("")
    documento.add_paragraph("")
    documento.add_paragraph(f"     El Departamento de Escuelas notifica una multa por Planificacion de manera formal.")
    documento.add_paragraph("")
    documento.add_paragraph(
        f"     El instructor {nombre.capitalize()} no entrego la planificacion de: {escuelas_formateadas}. Es por tales motivo asistio sin planificacion alguna. Esto se debió a que no entregó la información solicitada oportunamente al Departamento de Escuelas el dia Domingo {fecha_formateada}."
    )
    documento.add_paragraph("")
    documento.add_paragraph("Att: Ing. Daniel Cazorla")

    if not os.path.exists("multas"):
        os.makedirs("multas")
        documento.save(f"multas/multa-Planificacion-{nombre}-{fecha_formateada}.docx")

    documento.save(f"multas/multa-Planificacion-{nombre}-{fecha_formateada}.docx")
    return "multas/"

escuelas = []
while True:
    while True:
        escuela=input("Nombre de escuela: ").lower()
        escuelas.append(escuela)

        q = input("Otra escuela? [y][n] ").lower()
        if q == 'n':
            # pdf('multas/')
            break

    print("="*20)
    ruta = mp(
        nombre=input("nombre: ").lower(),
        escuela=escuelas,
        dia="domingo"
    )

    q = input("continuar [y][n] ").lower()
    if q == 'n':
        word_to_pdf(ruta)
        break